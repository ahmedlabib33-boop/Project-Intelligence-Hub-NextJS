"""Generate the Project Intelligence Hub technical and board handover package.

This is deliberately self-contained so the handover can be regenerated after
new projects, tests, routes, or operating scripts are added.
"""

from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CANONICAL_ROOT = Path(r"D:\one drive data\OneDrive\Documents\Project Intelligence Hub")
OUT = Path(__file__).resolve().parent
NOW = datetime.now().strftime("%d %B %Y, %H:%M")

NAVY = "0B2545"
BLUE = "2E74B5"
TEAL = "0F7C82"
GOLD = "B8821E"
INK = "1B263B"
MUTED = "617083"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
PALE_TEAL = "E7F5F4"
PALE_GOLD = "FFF4DF"
WHITE = "FFFFFF"


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, color: str = "D7DEE8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_fixed_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_run(run, *, name="Calibri", size=11, color=INK, bold=False, italic=False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph(paragraph, before=0, after=6, line=1.1, alignment=None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def configure_document(doc: Document, landscape=False) -> None:
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        code.font.size = Pt(8.5)
        code.font.color.rgb = RGBColor.from_string(WHITE)
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(8)
        code.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.text = "SAMCO EGYPT | PROJECT INTELLIGENCE HUB"
    set_paragraph(header, 0, 0, 1.0)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run(run, size=8, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.text = "Designed & Created | "
    set_paragraph(footer, 0, 0, 1.0)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base = footer.runs[0]
    set_run(base, size=8.5, color=MUTED)
    named = footer.add_run("Engr. Ahmed Labib")
    set_run(named, size=8.5, color=TEAL, bold=True)
    footer.add_run(" | Internal Handover")
    set_run(footer.runs[-1], size=8.5, color=MUTED)


def add_title(doc: Document, title: str, subtitle: str, *, label: str) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, 18, 2, 1.0)
    r = p.add_run(label.upper())
    set_run(r, size=10, color=TEAL, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, 0, 4, 1.0)
    r = p.add_run(title)
    set_run(r, size=27, color=NAVY, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, 0, 18, 1.15)
    r = p.add_run(subtitle)
    set_run(r, size=13, color=MUTED)
    rule = doc.add_paragraph()
    set_paragraph(rule, 0, 12, 1.0)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), TEAL)
    borders.append(bottom)
    p_pr.append(borders)


def add_cover_metadata(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    set_fixed_table(table)
    widths = [Inches(1.6), Inches(4.9)]
    for label, value in rows:
        cells = table.add_row().cells
        for i, cell in enumerate(cells):
            cell.width = widths[i]
            set_cell_margin(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_fill(cells[0], LIGHT_BLUE)
        p = cells[0].paragraphs[0]
        set_paragraph(p, 0, 0, 1.0)
        set_run(p.add_run(label), size=9, color=NAVY, bold=True)
        p = cells[1].paragraphs[0]
        set_paragraph(p, 0, 0, 1.0)
        set_run(p.add_run(value), size=9.5, color=INK)


def add_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None, style=None) -> None:
    p = doc.add_paragraph(style=style)
    set_paragraph(p)
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix):]))
    else:
        set_run(p.add_run(text))


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph(p, 0, 3, 1.12)
        for run in p.runs:
            set_run(run, size=10.25)
        if not p.runs:
            set_run(p.add_run(item), size=10.25)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph(p, 0, 4, 1.12)
        if p.runs:
            for run in p.runs:
                set_run(run, size=10.25)
        else:
            set_run(p.add_run(item), size=10.25)


def add_callout(doc: Document, label: str, text: str, tone: str = "teal") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_fixed_table(table)
    cell = table.cell(0, 0)
    set_cell_margin(cell, top=130, start=160, bottom=130, end=160)
    set_cell_border(cell, TEAL if tone == "teal" else GOLD)
    set_cell_fill(cell, PALE_TEAL if tone == "teal" else PALE_GOLD)
    p = cell.paragraphs[0]
    set_paragraph(p, 0, 0, 1.1)
    set_run(p.add_run(f"{label}: "), size=10, color=NAVY, bold=True)
    set_run(p.add_run(text), size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc: Document, text: str) -> None:
    for line in text.strip().splitlines():
        p = doc.add_paragraph(style="Code Block")
        set_paragraph(p, 0, 0, 1.0)
        p.paragraph_format.left_indent = Inches(0.12)
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), NAVY)
        p._p.get_or_add_pPr().append(shade)
        set_run(p.add_run(line), name="Consolas", size=8.3, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_fixed_table(table)
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        set_cell_margin(cell)
        set_cell_fill(cell, LIGHT_BLUE)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph(p, 0, 0, 1.0)
        set_run(p.add_run(header), size=8.7, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row in rows:
        table_row = table.add_row()
        prevent_row_split(table_row)
        cells = table_row.cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.width = Inches(widths[index])
            set_cell_margin(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_paragraph(p, 0, 0, 1.0)
            set_run(p.add_run(str(value)), size=8.7, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def line_count(path: Path) -> int:
    try:
        return sum(1 for _ in path.open(encoding="utf-8", errors="ignore"))
    except OSError:
        return 0


def read_json(path: Path) -> dict:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}


def project_records() -> list[dict]:
    records = []
    for manifest in sorted(ROOT.glob("projects/**/project_manifest.json")):
        data = read_json(manifest)
        if not data:
            continue
        if manifest.parent.name == "_PROJECT_TEMPLATE":
            continue
        records.append({
            "sector": manifest.parent.parent.name,
            "folder": manifest.parent.name,
            "project_id": data.get("project_id", "Missing"),
            "display": data.get("project_display_name", manifest.parent.name),
            "status": data.get("status", "Active"),
        })
    return records


def code_files() -> list[Path]:
    roots = [ROOT / "src", ROOT / "tools", ROOT / "website" / "src", ROOT / "tests"]
    extensions = {".py", ".ts", ".tsx", ".ps1", ".bat", ".json", ".md"}
    files: list[Path] = []
    for root in roots:
        if not root.exists():
            continue
        for path in root.rglob("*"):
            if not path.is_file() or path.suffix.lower() not in extensions:
                continue
            if any(part in {"node_modules", ".next", "__pycache__", ".venv", ".venv-analytics"} for part in path.parts):
                continue
            files.append(path)
    for path in [
        ROOT / "dashboard.py",
        ROOT / "contract_claims_center.py",
        ROOT / "RUN_FULL_PROJECT_NO_GIT_SYNC.bat",
        ROOT / "RUN_VERCEL_PIPELINE_TEST.bat",
        ROOT / "VERCEL_NEXTJS_DEPLOYMENT.md",
    ]:
        if path.exists() and path not in files:
            files.append(path)
    return sorted(files, key=lambda item: item.as_posix().lower())


def classify_code_file(path: Path) -> str:
    name = path.name.lower()
    parts = {part.lower() for part in path.parts}
    if name == "dashboard.py":
        return "Streamlit shell, dashboard tabs, project selector, metric composition, local exports"
    if name == "contract_claims_center.py":
        return "Project-local contract extraction, clause library, claim evidence, draft and rebuttal workflow"
    if "steel_delay_tia" in name:
        return "Canonical TIA calculation: data audit, event mapping, relationship logic, fragnet and concurrency assessment"
    if "project_context" in name or "project_catalog" in name:
        return "Project discovery, manifest identity, folder resolution and project-id isolation"
    if "letters" in name:
        return "Project letter ingest, matching, thread/risk analysis and export support"
    if "guardrail" in name:
        return "Data-quality warnings and publish guardrail reporting"
    if "generate_nextjs" in name:
        return "Builds selected-project Vercel payloads and static output inventory from canonical data"
    if "vercel_project_pipeline" in name:
        return "Validated generate, test, GitHub publish, Vercel deploy and watcher workflow"
    if "github_no_git" in name or "sync" in name:
        return "Repository synchronization and local state management without Git CLI"
    if "report" in name or "output" in name:
        return "Output Studio report, artifact manifest or delivery helper"
    if "route.ts" in name or "api" in parts:
        return "Server-side API endpoint with project scope and controlled error handling"
    if "page.tsx" in name:
        return "Next.js same-page dashboard and selected-project workspace user interface"
    if "component" in parts:
        return "Reusable website user-interface component"
    if "test" in name or "tests" in parts:
        return "Automated regression and isolation test"
    return "Supporting application, configuration, documentation or validation component"


def test_inventory() -> list[list[str]]:
    rows: list[list[str]] = []
    for path in sorted((ROOT / "tests").glob("test_*.py")):
        functions = re.findall(r"^def (test_[A-Za-z0-9_]+)", path.read_text(encoding="utf-8", errors="ignore"), flags=re.M)
        summary = "; ".join(item.replace("test_", "").replace("_", " ") for item in functions[:3])
        if len(functions) > 3:
            summary += f"; plus {len(functions) - 3} more"
        rows.append([path.name, str(len(functions)), summary or "Module-level validation"])
    return rows


def add_toc(doc: Document, headings: list[str]) -> None:
    doc.add_heading("Contents", level=1)
    for index, heading in enumerate(headings, start=1):
        p = doc.add_paragraph()
        set_paragraph(p, 0, 2, 1.05)
        set_run(p.add_run(f"{index:02d}. {heading}"), size=10.2, color=NAVY)
    doc.add_page_break()


def make_handover() -> Path:
    doc = Document()
    configure_document(doc)
    add_title(
        doc,
        "Complete Application Handover",
        "Project Intelligence Hub | Canonical Streamlit operations, Next.js/Vercel delivery, data lineage, tests, watchers and maintenance runbook.",
        label="Technical Reference",
    )
    add_cover_metadata(doc, [
        ("Prepared for", "SAMCO Egypt management, planning, controls and technology teams"),
        ("System", "Project Intelligence Hub"),
        ("Canonical calculation layer", str(CANONICAL_ROOT)),
        ("Hosted delivery layer", str(ROOT / "website")),
        ("Public site", "https://samcoegyptdashboard.vercel.app"),
        ("Generated", NOW),
        ("Document purpose", "Operate, verify, extend and hand over the system without relying on the original developer"),
    ])
    add_callout(doc, "How to use this document", "Start with Sections 1 to 5 for operational understanding. Use Sections 8 to 11 for running, testing, publishing and troubleshooting. Appendix A is the code inventory.")
    doc.add_page_break()

    toc = [
        "Executive Summary", "Operating Model", "Architecture and Data Flow", "Project Identity and Isolation",
        "Project Folder Standard", "Feature Catalogue", "Data Lineage", "Pipeline and Watchers",
        "Run, Test and Publish Commands", "VS Code and Codex Prompt Library", "Troubleshooting and Maintenance",
        "Release Checklist", "Appendix A: Code and Test Inventory",
    ]
    add_toc(doc, toc)

    doc.add_heading("1. Executive Summary", level=1)
    add_paragraph(doc, "Project Intelligence Hub is a project-controls operating system. It reads separate project files, validates and calculates project controls information locally, then publishes a project-scoped view to the Vercel website.")
    add_paragraph(doc, "There are two connected layers. The Streamlit/Python layer is the calculation and controlled document-ingestion layer. The Next.js/Vercel layer is the mobile-friendly management and selected-project viewing layer. The Vercel site does not write to local project source files; its Universal Report Engine creates an editable PowerPoint only in the user's browser from the active project payload and the current form/CSV inputs.")
    add_callout(doc, "Non-negotiable control", "project_id is the identity boundary. Every project payload, output artifact, TIA result, claims record and AI context must carry project_id and project_key. A missing selected-project file produces an empty state; it must never use a different project's data.")
    add_table(doc, ["Layer", "Plain-English purpose", "Main operator"], [
        ["Local project folders", "The controlled source of truth for project data, contracts, letters, schedules and evidence.", "Planning / project team"],
        ["Python Streamlit engines", "Read source files, calculate project controls, TIA and claims intelligence, and build local outputs.", "Planning / claims / controls"],
        ["Generated payloads", "Project-isolated JSON, report manifests and static artifacts produced from the local calculation layer.", "Pipeline"],
        ["Next.js / Vercel", "Management dashboard and selected-project workspace with direct report downloads and optional AI assistance.", "Board / management / project leads"],
    ], [1.35, 3.85, 1.3])

    doc.add_heading("2. Operating Model", level=1)
    add_paragraph(doc, "The application should be understood as a controlled chain, not as a single web page.")
    add_code(doc, "Project folders -> Streamlit / Python engines -> generated project payloads and artifacts -> validation -> GitHub sync -> Vercel deployment -> management view")
    add_numbered(doc, [
        "A project team adds or updates files inside its own project folder.",
        "The local engine discovers the project through its manifest, applies project_context and calculates only that project's results.",
        "The website generator exports the selected project's JSON, data confidence, source lineage and output manifest.",
        "The validator checks project identity, available feature blocks, report artifacts and no cross-project fallback.",
        "The no-Git sync helper publishes the validated workspace through the GitHub API. Vercel builds the website from the published repository.",
        "The Vercel watcher watches canonical project files and website code every 30 seconds, then repeats the validated release path when content changes.",
    ])
    add_callout(doc, "Important limitation", "Browser users can review, search, complete the Universal Report Engine form, load a report-input CSV and download an editable PowerPoint. The CSV is used only for that browser report and is not written to project storage. P6 ingestion, contract/letter rebuilding and TIA recalculation remain controlled local operations.", tone="gold")

    doc.add_heading("3. Architecture and Data Flow", level=1)
    add_table(doc, ["Area", "Canonical files / folders", "What it does", "Output surface"], [
        ["Project discovery", "projects/<sector>/<project>/project_manifest.json", "Assigns stable project_id and display information; supports folder rename and sector grouping.", "Dropdown, dashboard filters, project workspace"],
        ["Controls data", "01-data and 03-schedule", "Provides projects, activities, WBS, milestones, EVM, costs, risks and S-curves.", "Overview, WBS, Activities, Milestones, S-Curve, EVM, Risks"],
        ["Delay / TIA", "02-delay_analysis", "Uses templates, P6 export, relationship file, supply/IFC/RFI/payment and concurrency evidence.", "Delay Analysis - Time Impact Analysis"],
        ["Contract / evidence", "05-contracts and 06-evidence", "Builds clause libraries, evidence maps, entitlement and claim support.", "Contracts and Contract & Claims Intelligence Center"],
        ["Letters", "07-letters_intelligence/inbox", "Ingests, classifies, links and groups correspondence threads by project.", "Letters Intelligence"],
        ["Brand and reports", "08-branding, 10-deliverables, 11-outputs", "Applies project branding and stores generated report artifacts.", "Output Studio and direct downloads"],
    ], [1.1, 1.7, 2.25, 1.45])
    add_paragraph(doc, "The Vercel generator remains a publisher, not a second calculation engine. It consumes locally calculated project data and report artifacts so cost, progress, EVM and TIA logic are not recalculated differently in browser JavaScript. The Universal Report Engine presentation builder is a packaging surface only: it places active-project values and user-supplied report rows into the approved presentation format without creating analytical conclusions.")

    doc.add_heading("4. Project Identity and Isolation", level=1)
    add_bullets(doc, [
        "Every active project has a project_manifest.json with project_id, project_key, display name, folder name and operational status.",
        "Sector is determined by the parent folder. Adding a new sector and project folder is supported without project-specific source-code registration.",
        "Folder renaming updates the display/folder path while the manifest preserves project_id, avoiding duplicate identity where a manifest remains with the folder.",
        "All project paths are resolved through project_context. Loaders must receive context and search only the selected project's folders.",
        "Portfolio mode aggregates explicit project records. It must label its results as portfolio-level and must retain each contributing project_id.",
        "A missing source file is a controlled setup state. It is not permission to show data from a previous project, global root file or default project.",
    ])
    add_table(doc, ["Current project", "Sector", "project_id", "Source isolation rule"], [
        [row["display"], row["sector"], row["project_id"], "Reads only its own manifest, data, evidence and output path"]
        for row in project_records()
    ], [2.45, 1.05, 1.85, 1.15])

    doc.add_heading("5. Project Folder Standard", level=1)
    add_paragraph(doc, "Each project follows the same numbered structure. Only populate a folder when the project actually owns that information; do not copy records between projects simply to fill an empty module.")
    add_table(doc, ["Folder", "Put this here", "Used by"], [
        ["01-data", "Core project control CSV/XLSX: activities, cost, EVM, risk, milestones, progress and import templates.", "Overview, WBS, Activities, EVM, Risks, Decision Dashboard"],
        ["02-delay_analysis", "TIA methodology, template CSVs, P6 activity export, relationship file, supply, IFC/RFI/payment, concurrency and MEP evidence.", "Delay Analysis - Time Impact Analysis"],
        ["03-schedule", "Baseline/current schedule support files, WBS and schedule references.", "WBS, milestones, S-Curve, schedule comparison"],
        ["04-source_excel", "Original received Excel workbooks retained as source evidence.", "Traceability and controlled imports"],
        ["05-contracts", "Contract source files, clauses and the project-local clause library/knowledge base.", "Contracts and Claims Intelligence"],
        ["06-evidence", "Evidence register, photos, document references and claim/TIA supporting material.", "Claims, TIA and report appendices"],
        ["07-letters_intelligence", "Letter workbooks and inbox folders for contractor/consultant correspondence.", "Letters Intelligence and claim evidence"],
        ["08-branding", "Logo, project identity and approved report branding.", "Headers and Output Studio"],
        ["09-notes", "Meeting, engineering and claims notes in project scope.", "Technical Advisor and review context"],
        ["10-deliverables", "Controlled deliverables generated or received for that project.", "Report and submission traceability"],
        ["11-outputs", "Auto-generated project reports and manifest; do not manually mix other project outputs here.", "Output Studio and Vercel downloads"],
        ["12-logs", "Project-specific operational logs where enabled; shared sync logs live at workspace 12-logs.", "Troubleshooting"],
    ], [1.25, 3.5, 1.75])

    doc.add_heading("6. Feature Catalogue", level=1)
    add_paragraph(doc, "The Vercel Decision Making Dashboard remains the portfolio command view. When a project is selected from the same-page selector, only one selected-project workspace tab is visible at a time. The design is shared, but the data scope changes to the active project.")
    add_table(doc, ["Workspace area", "What the user sees", "Main data source / calculation"], [
        ["Decision Making Dashboard", "Portfolio KPIs, management decision brief, early warnings, sectors, scenarios, action tracking and AI search.", "All generated project records, filtered and explicitly labelled as portfolio scope"],
        ["Overview", "Identity, contract value, dates, progress, activity count, alerts and source confidence.", "Project metadata, activities, cost/EVM, letters and source status"],
        ["WBS and Activities", "Hierarchy, critical/deviated/RFT activities, float, forecast slip and progress variance.", "Project WBS/activity exports"],
        ["Milestones and S-Curve", "Milestone status, change-order exposure, planned/actual/invoiced progress.", "Milestones, S-Curve and contract data"],
        ["EVM Analysis", "BAC, PV, EV, AC, SPI, CPI, CV, SV, EAC and TCPI with root cause context.", "Project EVM and commercial source files; no invented values"],
        ["Analytics Intelligence", "Warnings, confidence, scenario calculations and AI-based explanation.", "Derived source-backed indicators plus optional AI"],
        ["Contracts", "Value, certification, paid/unpaid, retention and variations.", "Project contracts, certificates and payment records"],
        ["Letters Intelligence", "Inbox detector, registers, linked correspondence, issue threads, alerts and AI review.", "07-letters_intelligence only"],
        ["Risks", "Risk register, exposure, owners, response and IFC/RFI links.", "Risk register plus linked source evidence"],
        ["Delay Analysis - TIA", "Six subviews: Uploads, Tables & Conclusion, MEP Activities, AI - TIA, Question and Download Reports.", "02-delay_analysis templates and canonical TIA engine"],
        ["Contract & Claims Intelligence", "Clauses, evidence mapping, rebuttals, claim builder and project-scoped exports.", "05-contracts, 06-evidence and project claims database"],
        ["Technical Advisor", "Question-bank guidance separated from confirmed project evidence.", "Selected project payload plus shared technical knowledge bank"],
        ["Conference", "Project meeting URL and inline review context.", "Project meeting configuration only"],
        ["Output Studio", "Executive, master, elite SVG and linked dashboards, plus a form/CSV-driven Universal Report Engine that downloads an editable PowerPoint.", "Active ProjectRecord, user form/CSV input and project output manifest"],
    ], [1.55, 2.7, 2.25])

    doc.add_page_break()
    doc.add_heading("Universal Project Report Engine - production update", level=2)
    add_paragraph(doc, "Production release 4f58587 was published to GitHub main and deployed by the connected Vercel production project on 18 August 2026. The public alias is https://samcoegyptdashboard.vercel.app.")
    add_table(doc, ["Control", "Implemented behavior", "Verification evidence"], [
        ["Report selection", "Uses the selected family from the existing 30-family engine catalogue; the UI does not maintain a separate report list.", "Public Output Studio catalogue and production form verified"],
        ["Project inputs", "Prefills project name, project_id, key, sector, progress, SPI, CPI, delay exposure, risk and data quality from the active ProjectRecord.", "Switch test replaced ROYA values with Sophia values and retained no ROYA CSV values"],
        ["User inputs", "Accepts reporting period, author, status, notes and editable/uploaded Metric, Value, Status and Notes CSV rows.", "CSV template download and completed CSV workflow verified"],
        ["Download", "Generates a three-slide editable PPTX in the browser: cover, report table/KPIs and input/governance register.", "Production download completed with no browser console errors"],
        ["Isolation", "The component receives only project={project}, performs no fetch and resets form/CSV state when the active project changes.", "18 targeted tests passed; 115 pipeline checks passed; deck contained only the selected project identity"],
    ], [1.15, 3.45, 2.15])
    add_callout(doc, "Report governance", "A browser-generated report is a presentation of supplied inputs. It does not replace controlled local evidence validation, native schedule calculations, contractual determination or release approval.", tone="gold")

    doc.add_heading("7. Delay Analysis - Time Impact Analysis", level=1)
    add_paragraph(doc, "The TIA module is evidence-led. It does not fabricate a delay, grant EOT or treat compensation as automatic. Its function is to show what can safely be modelled, what is supporting evidence, and what must be checked in Primavera P6.")
    add_table(doc, ["TIA view", "Purpose", "Key evidence"], [
        ["Uploads", "Shows recognized data and lets the reviewer include/exclude scenario support streams without altering source evidence.", "Required templates, P6 and evidence file inventory"],
        ["Tables & Conclusion", "Provides data audit, event register, mapping, fragnet recommendation, concurrency and entitlement evidence.", "Canonical TIA results and selected project source tables"],
        ["MEP Activities", "Shows MEP activity schedule and civil interface logic with linked correspondence context.", "MEP activity, MEP schedule, civil logic and letters"],
        ["AI - TIA", "Explains method selection, dependencies, fragnet and relationship evidence in structured form.", "Canonical engine and TIA methodology / submitted guide where project-specific"],
        ["Question", "Answers a focused question with project-only evidence and safe AI guidance.", "Event, RFI, IFC, concurrency and canonical analysis tables"],
        ["Download Reports", "Offers the selected project's TIA-supporting reports and supporting exports.", "Project output manifest and direct mobile-safe links"],
    ], [1.25, 3.2, 2.05])
    add_bullets(doc, [
        "05-relationship_file.csv is active evidence: it normalizes predecessor/successor IDs, FS/SS/FF/SF relationships and lags, then attaches driving logic to impacted activities.",
        "Employer steel, IFC, RFI, payment and contractor mitigation records are treated differently according to evidence and contract position. Payment remains commercial/supporting evidence unless a direct schedule/resource link is shown.",
        "P6 critical path, float, fragnet and concurrency must be checked before a final EOT conclusion. Non-recalculated results are visibly marked indicative.",
    ])

    doc.add_heading("8. Data Lineage", level=1)
    add_paragraph(doc, "Use data_to_program.md and data_lineage.json as the detailed source-to-screen register. The summary below explains the normal reading path for management users.")
    add_table(doc, ["Displayed result", "Where it originates", "Transformation / control"], [
        ["Project title, sector, folder", "project_manifest.json and parent sector folder", "Project discovery preserves project_id; folder name is presentation data"],
        ["Contract value, paid and remaining", "Project controls / contracts / payment data", "Selected project aggregation only; invalid numeric values are shown as N/A"],
        ["Progress, planned progress, SPI/CPI", "Progress and EVM source tables", "EVM formula requires valid BAC/PV/EV/AC; denominators of zero are not used"],
        ["Risk, claims and decision alerts", "Project risk, claims, letters, contracts and guardrail findings", "Threshold and evidence-driven; missing evidence is visible"],
        ["Letters status", "07-letters_intelligence inbox and register", "Automatic direction/reference/thread detection stays inside selected project"],
        ["TIA conclusion", "Project 02-delay_analysis templates and canonical engine", "P6/relationship/concurrency evidence is required; EOT remains indicative without recalculation"],
        ["Download files", "11-outputs/<project>/.report_manifest.json", "Artifact includes project identity, source fingerprint, file type and checksum"],
    ], [1.45, 2.2, 2.85])

    doc.add_heading("9. Pipeline and Watchers", level=1)
    add_paragraph(doc, "The watcher uses file content fingerprints, not directory timestamps alone. This avoids republishing due to ordinary folder timestamp changes while still detecting a new project, a renamed project folder, source-table changes and tracked website-code changes.")
    add_table(doc, ["Stage", "Tool", "Outcome", "Log / state"], [
        ["Discover and calculate", "tools/generate_nextjs_website_data.py", "Generates project JSON, source confidence, canonical TIA snapshot and output references.", "website/public/data and project output manifests"],
        ["Validate", "tools/validate_streamlit_vercel_pipeline.py", "Checks project identity, payload contract, TIA relation data and artifact presence.", "12-logs/vercel_streamlit_pipeline_audit_latest.md"],
        ["Build", "website npm run build", "Runs TypeScript/build validation before release.", "Next.js build output"],
        ["GitHub publish", "tools/github_no_git_sync.ps1", "Uses GITHUB_TOKEN or GH_TOKEN without Git CLI; secrets are excluded.", "12-logs/github_sync.log and .sync_state/local_manifest.json"],
        ["Vercel deploy", "tools/vercel_project_pipeline.ps1", "Deploys production build and validates public assets.", "12-logs/vercel_project_pipeline.log and pipeline state"],
        ["Watch", "vercel_project_pipeline.ps1 -Mode Watch", "Polls tracked source and site files every 30 seconds and releases only when content changes.", ".sync_state/vercel_project_pipeline_state.json"],
    ], [1.1, 1.8, 2.35, 1.25])
    add_callout(doc, "Credentials", "Never put API keys or GitHub tokens inside source files, reports or prompts. The sync helper reads only GITHUB_TOKEN or GH_TOKEN. AI providers read their own configured environment variables. .env.local and secrets are excluded from synchronization.", tone="gold")

    doc.add_heading("10. Run, Test and Publish Commands", level=1)
    add_callout(
        doc,
        "Latest verification recorded for this handover",
        "On 18 August 2026, the clean Universal Report Engine release passed the Next.js production build, 18 targeted isolation/report tests and 115 pipeline validation checks with zero failures. The generated PPTX rendered as three slides with no overflow; its XML contained the selected ROYA project name and ID and none of the four other project names. VERCEL_SYNC DryRun passed watcher detection, GitHub main advanced to 4f58587, and Vercel deployment dpl_6cKrnvb2p23vrS7hTwjgGUorBG9u reached Ready with the production alias attached.",
        tone="teal",
    )
    add_paragraph(doc, "Run commands in PowerShell from the stated directory. Do not run deployment commands while an unrelated long-running deployment is active. Commands below deliberately do not include secret values.")
    doc.add_heading("Canonical Streamlit operations", level=2)
    add_code(doc, rf'''
Set-Location "{CANONICAL_ROOT}"
& .\.venv\Scripts\python.exe -m pytest -q tests -p no:cacheprovider
& .\.venv\Scripts\streamlit.exe run dashboard.py
.\RUN_APP.bat
python tools\validate_project_isolation.py
''')
    doc.add_heading("Generate and validate Vercel data", level=2)
    add_code(doc, rf'''
Set-Location "{ROOT}"
& .\.venv-analytics\Scripts\python.exe tools\generate_nextjs_website_data.py
& .\.venv-analytics\Scripts\python.exe tools\validate_streamlit_vercel_pipeline.py
& .\.venv-analytics\Scripts\python.exe -m pytest -q tests -p no:cacheprovider

Set-Location "{ROOT}\website"
npm run dev
npm run build
''')
    doc.add_heading("Safe sync and deployment", level=2)
    add_code(doc, rf'''
Set-Location "{ROOT}"
# Inspect only; does not publish.
powershell.exe -NoProfile -ExecutionPolicy Bypass -File .\tools\vercel_project_pipeline.ps1 -Mode DryRun -IntervalSeconds 30

# One validated release: generate, validate, build, publish and deploy.
powershell.exe -NoProfile -ExecutionPolicy Bypass -File .\tools\vercel_project_pipeline.ps1 -Mode Once -IntervalSeconds 30

# Continuous watcher: checks for tracked content changes every 30 seconds.
powershell.exe -NoProfile -ExecutionPolicy Bypass -File .\tools\vercel_project_pipeline.ps1 -Mode Watch -IntervalSeconds 30

# Repository sync only; no Git CLI.
.\RUN_FULL_PROJECT_NO_GIT_SYNC.bat Once 30
''')

    doc.add_heading("11. VS Code and Codex Prompt Library", level=1)
    add_paragraph(doc, "Open the correct workspace before changing code. The canonical calculation workspace and the Vercel delivery workspace have different responsibilities.")
    add_code(doc, rf'''
code "{CANONICAL_ROOT}"
code "{ROOT}"
code "{ROOT}\website"
''')
    add_table(doc, ["Use case", "Safe prompt to give a coding assistant"], [
        ["Add a project", "Create a project folder under the correct sector. Create only missing standard folders and a project_manifest.json. Preserve existing files, assign a stable project_id, and do not copy data from another project."],
        ["Check isolation", "Trace this selected project from manifest through project_context, loaders, TIA, claims, letters, report outputs and generated Vercel JSON. Report every fallback path and fix only cross-project leakage."],
        ["Add a source field", "Add this field to the selected project data model, preserve source columns, include project_id in the result, document it in data_to_program.md, and add a regression test."],
        ["Update TIA", "Use only the selected project's 02-delay_analysis files. Do not fabricate dates or EOT. Apply P6 relationship logic, float and concurrency evidence; label results indicative until P6 recalculation."],
        ["Update claims", "Use only the selected project's 05-contracts, 06-evidence and claims database. Cite source files and rows; show missing evidence rather than copying any other project library."],
        ["Deploy safely", "Run generator, validator, tests and production build. Publish only after all pass, then verify public JSON and an HTML/PDF/PPTX artifact using the selected project's project_id."],
        ["Improve UI", "Preserve data flow, project selector and active project scope. Change only the requested component. Verify desktop and mobile without horizontal overflow and keep one tab panel visible at a time."],
    ], [1.25, 5.4])

    doc.add_heading("12. Troubleshooting and Maintenance", level=1)
    add_table(doc, ["Symptom", "Likely cause", "Check and corrective action"], [
        ["A project shows another project's information", "A loader, cache key or output path omitted project_id / project_context.", "Run isolation validator; inspect loader search order; clear selected project cache; confirm all SQL/JSON filters use active project_id."],
        ["EVM shows zero or N/A", "Missing or invalid BAC/PV/EV/AC, or percentage not normalized.", "Check source EVM/cost files. Do not insert zero to make a card look complete; correct the input and regenerate."],
        ["TIA has no conclusion", "Required P6, relationship, event or date fields are missing.", "Use Uploads and Tables & Conclusion data audit. Supply missing files; do not create an EOT conclusion from a narrative alone."],
        ["Claims look identical across projects", "Shared fallback clause library or database path is being used.", "Verify 05-contracts, 06-evidence and contract_claims.db resolve under the selected project folder."],
        ["Site did not update", "Generator, GitHub sync, Vercel deploy or public propagation did not complete.", "Read 12-logs/vercel_project_pipeline.log; run DryRun then Once; verify public project JSON and report URLs."],
        ["Sync authentication fails", "No valid local GITHUB_TOKEN/GH_TOKEN or insufficient repository Contents write permission.", "Configure a repository-scoped token in the local environment; never paste it into code or documents."],
        ["Mobile report downloads fail", "A JavaScript-only download was used or artifact was missing.", "Check report manifest and direct anchor URL. HTML/PDF/PPTX must be public static files with valid MIME type."],
    ], [1.4, 2.1, 3.15])

    doc.add_heading("13. Release Checklist", level=1)
    add_numbered(doc, [
        "Confirm the selected project has a valid project_manifest.json and its files are inside its own project folder.",
        "Run project-isolation and module tests before generating the website payload.",
        "Run the Vercel generator, then inspect generated project JSON for project_id, project_key, source status and feature blocks.",
        "Verify each project report manifest contains only that project's HTML, PDF and PPTX artifacts.",
        "Run npm run build from website.",
        "Use the validated pipeline Once mode, then confirm the public URL and sample downloads return 200.",
        "In Output Studio > Universal Report Engine - ML, confirm the form shows the active project_id, download the CSV template, generate a PPTX, and verify its governance register contains only the selected project identity.",
        "Review guardrail warnings. Warnings are management-visible; do not silently suppress missing or suspicious data.",
        "Keep the watcher running only on a controlled local machine that has the canonical project folders and required credentials.",
    ])

    doc.add_heading("Appendix A. Code and Test Inventory", level=1)
    add_paragraph(doc, "This appendix is generated from the current delivery workspace. The inventory is intentionally practical: it gives the next maintainer a starting point and a plain-English role for each implementation file.")
    inventory_rows = []
    for path in code_files():
        rel = path.relative_to(ROOT).as_posix()
        inventory_rows.append([rel, str(line_count(path)), classify_code_file(path)])
    for start in range(0, len(inventory_rows), 45):
        if start:
            doc.add_heading("Appendix A continued", level=2)
        add_table(doc, ["File", "Lines", "Role"], inventory_rows[start:start + 45], [2.55, 0.45, 3.65])

    doc.add_heading("Appendix B. Automated Test Inventory", level=1)
    add_table(doc, ["Test file", "Tests", "Coverage focus"], test_inventory(), [2.1, 0.55, 4.0])
    add_paragraph(doc, "Document scope: this report intentionally does not contain API keys, GitHub tokens, private endpoints or source-document contents. Refer to controlled local project folders for the underlying evidence.")

    out = OUT / "Project_Intelligence_Hub_Complete_Technical_Handover.docx"
    doc.save(out)
    return out


def make_board_guide() -> Path:
    doc = Document()
    configure_document(doc, landscape=True)
    add_title(
        doc,
        "Board Presentation Guide",
        "Plain-English narrative for presenting Project Intelligence Hub to board members, executive management and project directors.",
        label="Management Presentation",
    )
    add_cover_metadata(doc, [
        ("Audience", "Board members, executive management and project directors"),
        ("Recommended duration", "15 to 20 minutes, followed by live project demonstration"),
        ("Presenter", "Engr. Ahmed Labib | Planning Department"),
        ("Purpose", "Explain the controls value, evidence source, governance limits and rollout decision"),
        ("Generated", NOW),
    ])
    add_callout(doc, "Presentation rule", "Speak in simple business language. Show the live system only after explaining that every management number is traceable to the owning project folder and that missing evidence remains visible.")
    doc.add_page_break()

    slides = [
        ("1. The Management Problem", "Management receives schedules, payment status, claims, letters and risk information in different formats and at different times. This makes it hard to see what decision is needed now.", "Open the Decision Making Dashboard. Point to the portfolio KPIs and decision brief.", "The system does not replace project teams. It gives management one controlled view of information already owned by the projects."),
        ("2. What the Hub Does", "Project Intelligence Hub collects project-control information into one system. It shows the portfolio view first, then lets management enter a selected project workspace without opening a different product.", "Use the selector to move between Decision Making Dashboard and one project.", "The portfolio view is clearly portfolio-level. A selected project view uses only that project's evidence."),
        ("3. Data Is Kept Separate", "Each project has its own folder, identity and project_id. This prevents one project's contract, delay, letter or report from appearing in another project.", "Show the project folder standard or the active-project badge.", "When data is missing, the system shows a setup or data-quality message. It does not borrow information from another project."),
        ("4. Decision Making Dashboard", "The dashboard brings together contract value, paid value, progress, delay, risk, claims exposure, SPI, CPI and data confidence. It also highlights the decisions that need an owner.", "Show portfolio KPI cards, early warnings, sector analysis and management decision brief.", "The cards are not manual numbers. They are calculated from the active project records and refresh after the source files are updated."),
        ("5. Project Controls in One Place", "For a selected project, the workspace covers overview, WBS, activities, milestones, S-curve, EVM, contracts, letters, risks, Delay TIA, claims, technical advice, conference review and Output Studio.", "Open two or three representative tabs only. Keep one active panel visible at a time.", "Each tab is designed to answer a management question: where are we, what is at risk, what is the evidence, and what should happen next?"),
        ("6. Delay Analysis and Time Impact Analysis", "The Delay TIA module organizes P6 data, relationships, supply events, IFC issues, RFIs, payments, mitigation and concurrency. It distinguishes an observed delay from a proven entitlement.", "Show Tables & Conclusion, then the relationship logic / fragnet evidence.", "The system never promises EOT automatically. Non-recalculated results are marked indicative until Primavera P6 verification."),
        ("7. Contracts, Claims and Letters", "The Hub joins contract clauses, evidence folders and correspondence. It helps the team see notice risk, missing evidence, linked letters and a defendable claim position.", "Show Letters Intelligence issue threads and Contract & Claims evidence mapping.", "A claim is stronger when its event, clause, notice, evidence and schedule impact are connected. The system exposes gaps rather than hiding them."),
        ("8. AI as a Controlled Advisor", "AI can explain available project data, technical knowledge and evidence gaps. It is not allowed to invent cost, delay, dates, clauses or project facts.", "Open the AI panel or Technical Advisor. Ask an example question.", "AI answers separate confirmed project evidence, inference, missing evidence and general guidance. The selected-project AI context cannot use another project's data."),
        ("9. Reports and Outputs", "Output Studio prepares a consistent executive, master, elite chart and linked dashboard package. The same selected-project metric payload supports HTML, PDF and PowerPoint files.", "Show a report manifest or a direct download link.", "Direct download links are used alongside JavaScript support so reports work on desktop, Android and iOS."),
        ("10. Data Trust and Governance", "A data guardrail layer checks suspicious values, missing fields and incomplete evidence. It produces warnings for management and can block publishing if governance chooses that policy later.", "Show data confidence or guardrail status.", "Good management decisions require visible quality limits. The Hub makes data quality a decision input rather than an afterthought."),
        ("11. Automation and Publishing", "The local project folders remain the controlled source. When source content changes, the watcher generates project payloads, tests the system, publishes to GitHub and deploys Vercel.", "Show the pipeline diagram or log summary.", "This keeps the public management view aligned with local control data while avoiding uncontrolled browser uploads into project evidence folders."),
        ("12. Board Decision Requested", "Approve a phased operating model: standardize the project folders, appoint data owners, run a weekly controls review, and use the Hub as the common management evidence view.", "Return to the management decision brief.", "The immediate value is faster, evidence-based action. The longer-term value is a reusable control system for every future SAMCO project."),
    ]
    for index, (title, say, show, message) in enumerate(slides):
        if index:
            doc.add_page_break()
        doc.add_heading(title, level=1)
        add_callout(doc, "Say this", say)
        doc.add_heading("Show this in the live app", level=2)
        add_paragraph(doc, show)
        doc.add_heading("Key message for the board", level=2)
        add_paragraph(doc, message)
        doc.add_heading("Likely board question", level=2)
        questions = {
            0: "How will this improve decision speed?",
            1: "Is this a replacement for our existing systems?",
            2: "How do we prevent project data mixing?",
            3: "Can we trust the KPI values?",
            4: "What information is still missing?",
            5: "Does it calculate final EOT entitlement?",
            6: "Can the system help us defend a claim?",
            7: "Can AI make a wrong statement?",
            8: "Can reports be shared without desktop software?",
            9: "What happens when data is incomplete?",
            10: "Who controls publishing and deployment?",
            11: "What do you need from the board?",
        }[index]
        add_paragraph(doc, questions, bold_prefix=None)
        doc.add_heading("Recommended answer", level=2)
        answers = {
            0: "It replaces manual collection from multiple spreadsheets with a controlled view of decision triggers, evidence gaps and owners.",
            1: "No. It is a management intelligence layer over project controls, contracts, letters and schedules. It keeps source files in their controlled folders.",
            2: "project_id, manifest-based discovery and project_context bind every loader, output, report and AI context to one project. Automated tests check this.",
            3: "Values are source-backed and show data confidence. Invalid or missing values are shown as N/A or warnings, not silently corrected.",
            4: "The app identifies exactly which source file or evidence is missing, so the project team can close the correct gap.",
            5: "No. It prepares evidence, fragnet and concurrency analysis. Final EOT needs a verified Primavera P6 recalculation and contractual determination.",
            6: "It improves traceability and readiness by mapping events to clauses, letters, evidence and schedule impact. Legal/contractual approval still belongs to the responsible team.",
            7: "The AI is constrained to current app data and the shared knowledge bank. It must identify missing evidence and cannot present assumptions as confirmed project facts.",
            8: "Yes. The generated HTML, PDF and PowerPoint files have direct project-scoped download links designed for desktop and mobile browsers.",
            9: "The guardrail and module empty states expose the issue. Management can decide whether to collect the evidence, accept an assumption or defer a decision.",
            10: "The controlled local machine owns source ingestion and the watcher. The validated pipeline records its state and publishes only after build and validation checks pass.",
            11: "Approve the folder standard, nominate project data owners, and require weekly evidence and KPI validation before executive review.",
        }[index]
        add_paragraph(doc, answers)
    doc.add_page_break()
    doc.add_heading("Presenter Preparation Checklist", level=1)
    add_numbered(doc, [
        "Open the Decision Making Dashboard and select a project with complete representative data.",
        "Confirm the current public URL and one HTML/PDF/PowerPoint download before the meeting.",
        "Prepare one risk, one letter thread, one contract clause/evidence example and one TIA evidence example.",
        "Do not present an indicative TIA value as final EOT. Say what P6/evidence step is still required.",
        "Use data confidence openly. A visible evidence gap is a management control benefit, not a presentation failure.",
    ])
    out = OUT / "Project_Intelligence_Hub_Board_Presentation_Guide.docx"
    doc.save(out)
    return out


def make_runbook() -> Path:
    text = f"""# Project Intelligence Hub - Run, Test and Prompt Reference

Generated: {NOW}

## Workspace roles

- Canonical Streamlit calculation and ingestion workspace: `{CANONICAL_ROOT}`
- Vercel delivery workspace: `{ROOT}`
- Website folder: `{ROOT / 'website'}`
- Public site: `https://samcoegyptdashboard.vercel.app`

## Daily local checks

```powershell
Set-Location \"{CANONICAL_ROOT}\"
& .\\.venv\\Scripts\\python.exe -m pytest -q tests -p no:cacheprovider
python tools\\validate_project_isolation.py
```

## Generate Vercel data and test the delivery workspace

```powershell
Set-Location \"{ROOT}\"
& .\\.venv-analytics\\Scripts\\python.exe tools\\generate_nextjs_website_data.py
& .\\.venv-analytics\\Scripts\\python.exe tools\\validate_streamlit_vercel_pipeline.py
& .\\.venv-analytics\\Scripts\\python.exe -m pytest -q tests -p no:cacheprovider
Set-Location .\\website
npm run build
```

## Vercel pipeline

```powershell
Set-Location \"{ROOT}\"
# Check configuration only
powershell.exe -NoProfile -ExecutionPolicy Bypass -File .\\tools\\vercel_project_pipeline.ps1 -Mode DryRun -IntervalSeconds 30
# One validated release
powershell.exe -NoProfile -ExecutionPolicy Bypass -File .\\tools\\vercel_project_pipeline.ps1 -Mode Once -IntervalSeconds 30
# Continuous watcher
powershell.exe -NoProfile -ExecutionPolicy Bypass -File .\\tools\\vercel_project_pipeline.ps1 -Mode Watch -IntervalSeconds 30
```

## VS Code

```powershell
code \"{CANONICAL_ROOT}\"
code \"{ROOT}\"
code \"{ROOT / 'website'}\"
```

## Safe coding prompts

1. **Project isolation:** `Trace this selected project from manifest through project_context, loaders, TIA, claims, letters, report outputs and generated Vercel JSON. Report every fallback path and fix only cross-project leakage.`
2. **New project:** `Create a project folder under the correct sector. Create only missing standard folders and a project_manifest.json. Preserve existing files, assign a stable project_id, and do not copy data from another project.`
3. **TIA:** `Use only the selected project's 02-delay_analysis data. Do not fabricate dates or EOT. Use P6 relationship logic, float and concurrency evidence, then label unverified output indicative.`
4. **Claims:** `Use only the selected project's 05-contracts, 06-evidence and claims database. Cite source files and rows, and show missing evidence.`
5. **Deployment:** `Run generator, validator, tests and build. Publish only after all pass. Verify public project JSON and an HTML, PDF and PPTX report using the selected project identity.`

## Never do this

- Do not hardcode a project name as a fallback data source.
- Do not copy a contract library, letters register or output folder from one project to another.
- Do not place secrets in source, prompts, reports or synchronization configuration.
- Do not present a non-P6-recalculated TIA result as final EOT or compensation.
"""
    out = OUT / "RUN_TEST_AND_PROMPT_REFERENCE.md"
    out.write_text(text, encoding="utf-8")
    return out


def make_code_index() -> Path:
    lines = ["# Project Intelligence Hub - Code and Feature Index", "", f"Generated: {NOW}", "", "| File | Lines | Plain-English role |", "|---|---:|---|"]
    for path in code_files():
        lines.append(f"| `{path.relative_to(ROOT).as_posix()}` | {line_count(path)} | {classify_code_file(path)} |")
    out = OUT / "CODE_AND_FEATURE_INDEX.md"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    outputs = [make_handover(), make_board_guide(), make_runbook(), make_code_index()]
    for output in outputs:
        print(output)
