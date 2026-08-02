"""Build board and planner handover documents for the project-local chart release."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
CATALOG_PATH = ROOT / "config" / "chart_catalog.json"
BOARD_OUTPUT = REPORTS / "Project_Intelligence_Hub_Chart_Clone_Board_Report.docx"
PLANNER_OUTPUT = REPORTS / "Project_Intelligence_Hub_Chart_Clone_Planner_Handover.docx"
REGISTER_OUTPUT = REPORTS / "Project_Intelligence_Hub_Chart_Clone_Technical_Handover_Register.md"

NAVY = "0B1F36"
CYAN = "0EA5C6"
TEAL = "10B981"
AMBER = "D99414"
SLATE = "425466"
PALE = "E8F4F7"
PALE_GOLD = "FFF6DE"
WHITE = "FFFFFF"

VERCEL_TEMPLATES = [
    ("phase_progress.csv", "Phase-level progress history for executive phase and completion views."),
    ("discipline_progress_history.csv", "Dated planned, actual, and forecast progress by discipline."),
    ("activity_completion_history.csv", "Dated started and completed activity counts."),
    ("evm_period_history.csv", "Dated BAC, PV, EV, AC, SPI, and CPI where the standard EVM register has no period series."),
    ("planned_cash_flow.csv", "Planned cash out series for the cash-flow comparison; payment/certification remains in the normal payments source."),
    ("risk_assessment_history.csv", "Dated risk snapshots and before/after mitigation scores."),
    ("delay_event_classification.csv", "Verified delay-event classification linked to an existing event; it cannot create an event."),
    ("tia_recovery_scenario.csv", "P6 and evidence-backed baseline, impacted, and recovery comparison scenario."),
]


def _shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def _cell(cell, value: str, *, bold: bool = False, color: str = NAVY, size: int = 8.8) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        _shade(table.rows[0].cells[index], NAVY)
        _cell(table.rows[0].cells[index], header, bold=True, color=WHITE, size=8.5)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2 == 0:
                _shade(cells[index], PALE)
            _cell(cells[index], value, bold=index == 0)
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
    doc.add_paragraph()


def _configure(doc: Document, title: str, subtitle: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SAMCO PROJECT INTELLIGENCE HUB")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(CYAN)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Designed & Created | Engr. Ahmed Labib | Project Intelligence Hub")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(SLATE)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(92)
    cover.paragraph_format.space_after = Pt(10)
    run = cover.add_run(title)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(29)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(40)
    run = sub.add_run(subtitle)
    run.font.name = "Aptos"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(CYAN)
    callout = doc.add_table(rows=1, cols=3)
    for cell, label, value, color in zip(
        callout.rows[0].cells,
        ["Calculation source", "Data boundary", "Publication standard"],
        ["Canonical project engines", "project_id + project_key", "Evidence first"],
        [CYAN, TEAL, AMBER],
    ):
        _shade(cell, NAVY)
        _cell(cell, f"{label}\n{value}", bold=True, color=color, size=9)
    doc.add_page_break()


def _heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(15 if level == 1 else 9)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(17 if level == 1 else 12)
    run.font.color.rgb = RGBColor.from_string(CYAN if level == 1 else NAVY)


def _paragraph(doc: Document, text: str, *, emphasis: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.12
    if emphasis and text.startswith(emphasis):
        run = paragraph.add_run(emphasis)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
        text = text[len(emphasis):]
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(9.6)
    run.font.color.rgb = RGBColor.from_string("263548")


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Aptos"
        run.font.size = Pt(9.4)
        run.font.color.rgb = RGBColor.from_string("263548")


def _catalogue() -> list[dict]:
    payload = json.loads(CATALOG_PATH.read_text(encoding="utf-8"))
    return list(payload["charts"])


def _board_report(charts: list[dict]) -> None:
    doc = Document()
    _configure(doc, "CHART CLONE RELEASE", "Board Report: Project-Local Vercel Data, Intelligence, and Control")
    _heading(doc, "1. Executive Position")
    _paragraph(doc, "The project workspace now publishes a complete 36-slot chart contract for every selected project. The dark digital visual language, chart card composition, accent system, labels, legends, and tab ownership follow the approved chart reference. Values are never copied from the reference; they are calculated only from the selected project payload.")
    _table(doc, ["Board outcome", "How the release controls it", "Management benefit"], [
        ["No sample chart values", "Every chart slot has project source lineage and an awaiting-data state.", "Avoids decisions based on demonstration curves or false zero values."],
        ["Project isolation", "Rows are filtered by project_id and every payload carries project_key.", "Portfolio roll-up can be used without mixing project evidence."],
        ["Transparent readiness", "Missing inputs render a chart-sized readiness card naming the required project-local file.", "Data owners know exactly what to complete."],
        ["Controlled intelligence", "AI can explain evidence and gaps but cannot create delay, EOT, or claim conclusions.", "Keeps management summaries useful without overstating certainty."],
    ], [1.45, 3.1, 2.45])
    _heading(doc, "2. 36-Chart Coverage")
    grouped: dict[str, list[str]] = {}
    for chart in charts:
        grouped.setdefault(str(chart["tab"]), []).append(str(chart["title"]))
    _table(doc, ["Project tab", "Reference chart slots", "Data rule"], [
        [tab, "; ".join(titles), "Rendered from selected project data only; unavailable data remains a controlled readiness card."]
        for tab, titles in grouped.items()
    ], [1.3, 3.85, 1.85])
    _heading(doc, "3. TIA Governance")
    _paragraph(doc, "Delay Analysis - Time Impact Analysis is retained as internal project evidence and is not published in the public workspace until explicitly recalled. Its project payload remains governed by native schedule sources, relationship logic, float/critical-path evidence, verified event classification, and P6/evidence gates.")
    _table(doc, ["Control", "Rule"], [
        ["Event classification", "A verified classification must link to an existing project delay event and matching activity where stated."],
        ["Recovery comparison", "Draft scenarios are visible only as draft planning input; contractual use requires verified P6 and evidence references."],
        ["Relationship logic", "Predecessor, successor, relationship type, lag, float, longest path, and fragnet controls use project schedule evidence."],
        ["Entitlement", "No final EOT, compensation, concurrency, or critical delay conclusion without project evidence and required P6 verification."],
    ], [1.65, 5.35])
    _heading(doc, "4. Immediate Board Actions")
    _bullets(doc, [
        "Assign each project a data owner for the eight optional Vercel chart templates and the existing project-control registers.",
        "Require a monthly evidence and data-quality review before board information is used for recovery, commercial, or claim decisions.",
        "Approve a P6 verification protocol for any recovered TIA result before it is issued as a contractual conclusion.",
        "Use Output Studio for controlled HTML, PDF, PowerPoint, and supporting downloads; do not treat live workspace chart cards as client reports.",
    ])
    doc.save(BOARD_OUTPUT)


def _planner_guide(charts: list[dict]) -> None:
    doc = Document()
    _configure(doc, "PLANNER HANDOVER GUIDE", "Maintaining Project Data, Vercel Charts, Validation, and Publishing")
    _heading(doc, "1. Project Operating Rule")
    _paragraph(doc, "A project is identified by its manifest project_id and project_key. The folder may be renamed, but the project_id must remain stable. Never copy rows, reports, letters, contracts, or outputs from one project into another project folder.")
    _heading(doc, "2. Where to Update Data")
    _table(doc, ["Folder", "What belongs there", "What the app does"], [
        ["01-data/import_templates", "Project, WBS, activities, milestones, S-curve, EVM, contracts, payments, risks, and claims source tables.", "Provides the primary source pipeline for workspace metrics and charts."],
        ["02-delay_analysis", "P6/XER, relationships, event register, concurrency, evidence, and canonical TIA inputs.", "Builds the governed internal delay-analysis payload for that project only."],
        ["05-contracts and 06-evidence", "Authorized contract text, clauses, notices, evidence, and claim support.", "Feeds the project-local Contract & Claims and AI evidence context."],
        ["07-letters_intelligence", "Incoming/outgoing letters and correspondence registers.", "Feeds the project-local letters detector, thread, and evidence outputs."],
        ["vercel", "Only optional missing-data CSV templates listed below.", "Overrides the matching legacy source only when it has valid selected-project rows."],
    ], [1.45, 3.75, 1.8])
    _heading(doc, "3. Vercel Template Rules")
    _paragraph(doc, "The `vercel` folder is created automatically for every current and future project, including `_PROJECT_TEMPLATE`. Templates are header-only and non-destructive. Do not add sample rows. Every populated row must contain the owning project_id.")
    _table(doc, ["Template", "Use", "Validation and precedence"], [
        [name, purpose, "Valid Vercel rows win. A header-only file falls back to the normal project source. If both have rows, the app uses Vercel and reports a duplicate-source warning."]
        for name, purpose in VERCEL_TEMPLATES
    ], [1.65, 3.25, 2.1])
    _heading(doc, "4. Update and Publish Sequence")
    _table(doc, ["Step", "Planner action", "Expected control"], [
        ["1", "Update only the owning project folder and retain project_id on each imported row.", "No data is silently taken from another project."],
        ["2", "Run the local generator and validation before publishing.", "The generator refreshes selected-project JSON and output artifacts; the validator checks isolation and source contracts."],
        ["3", "Review chart readiness cards and the data-quality findings for that project.", "Missing input is visible as an action, never plotted as zero or substituted."],
        ["4", "Run the Vercel publisher/watch workflow after validation passes.", "The publisher regenerates, builds, synchronizes, deploys, and checks public assets."],
        ["5", "Use Output Studio to issue formal outputs.", "HTML, PDF, PowerPoint, and supporting files remain project-scoped and traceable."],
    ], [0.5, 3.9, 2.6])
    _heading(doc, "5. Chart-to-Source Checklist")
    _table(doc, ["Tab", "Chart source priority", "Missing-data message"], [
        [str(chart["tab"]), ", ".join(chart.get("sources", [])), "Same-sized readiness card naming the required source."]
        for chart in charts
    ], [1.25, 4.1, 1.65])
    _heading(doc, "6. Troubleshooting")
    _bullets(doc, [
        "A chart shows awaiting data: complete the named source for the selected project; do not copy a row from another project.",
        "A chart data warning appears: check project_id, date format, duplicate periods, event/activity references, and required verification status.",
        "A Vercel template does not appear: re-run project discovery/generation; it creates missing templates without overwriting user files.",
        "A TIA result is not final: check the P6 reference, relationship evidence, concurrency, authorized contract route, and evidence reference.",
        "A public site value is stale: run the generator, the validation script, and then the Vercel publisher. Confirm the project JSON timestamp before reviewing the website.",
    ])
    doc.save(PLANNER_OUTPUT)


def _technical_register(charts: list[dict]) -> None:
    lines = [
        "# Project Intelligence Hub - Chart Clone Technical Handover Register",
        "",
        "## Architecture",
        "`project folder -> canonical Python calculations -> generated project JSON -> same-page Next.js workspace -> Output Studio artifacts`",
        "",
        "## Source Precedence",
        "1. Populated `project/vercel/<template>.csv` rows matching `project_id`.",
        "2. Existing project-local source defined in the chart mapping.",
        "3. Controlled awaiting-data card. No sample values, cross-project fallback, or zero placeholder.",
        "",
        "## Chart Register",
        "",
        "| ID | Tab | Type | Primary sources | Required columns |",
        "|---|---|---|---|---|",
    ]
    for chart in charts:
        lines.append(
            f"| {chart['id']} | {chart['tab']} | {chart['type']} | {', '.join(chart.get('sources', []))} | {', '.join(chart.get('required_columns', []))} |"
        )
    lines.extend([
        "",
        "## Required Vercel Templates",
        "",
        "| File | Mandatory boundary | Purpose |",
        "|---|---|---|",
    ])
    for name, purpose in VERCEL_TEMPLATES:
        lines.append(f"| `vercel/{name}` | `project_id` | {purpose} |")
    lines.extend([
        "",
        "## Deployment and Validation",
        "",
        "- Generator: `D:\\Project Intelligence Hub NextJS\\tools\\generate_nextjs_website_data.py`",
        "- Chart payload builder: `D:\\Project Intelligence Hub NextJS\\tools\\project_chart_payloads.py`",
        "- Pipeline validator: `D:\\Project Intelligence Hub NextJS\\tools\\validate_streamlit_vercel_pipeline.py`",
        "- Full publisher: `D:\\Project Intelligence Hub NextJS\\tools\\vercel_project_pipeline.ps1`",
        "- Website chart renderer: `D:\\Project Intelligence Hub NextJS\\website\\src\\app\\page.tsx`",
        "- The dashboard only renders chart payloads matching the active project identity.",
        "- Delay Analysis - Time Impact Analysis remains internal unless explicitly recalled.",
    ])
    REGISTER_OUTPUT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)
    charts = _catalogue()
    if len(charts) != 36:
        raise SystemExit(f"Expected 36 chart definitions; found {len(charts)}.")
    _board_report(charts)
    _planner_guide(charts)
    _technical_register(charts)
    print(BOARD_OUTPUT)
    print(PLANNER_OUTPUT)
    print(REGISTER_OUTPUT)


if __name__ == "__main__":
    main()
